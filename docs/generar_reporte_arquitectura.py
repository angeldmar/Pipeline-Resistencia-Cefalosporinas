"""Genera el reporte de "Analisis de resultados" sobre la arquitectura y el
producto final del sistema, en Word (.docx) y Markdown (.md).

Complementa al reporte estadistico (validacion_estadistica/): aquel analiza
la concordancia analitica; este describe el sistema construido —su
arquitectura, las decisiones de diseno y los productos finales (interfaz web,
reporte por muestra, contenedor)— con el mismo registro de una seccion de
analisis de resultados de tesis. El contenido (prosa y tablas) se define una
sola vez y se renderiza a ambos formatos. Las capturas de la interfaz ya
estan en docs/figuras/.

Uso:
    python docs/generar_reporte_arquitectura.py
Salida:
    docs/reporte_arquitectura.md        (navegable en el repositorio)
    docs/reporte_arquitectura.docx      (para la tesis)
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parent
FIGURES_DIR = DOCS_DIR / "figuras"
MD_PATH = DOCS_DIR / "reporte_arquitectura.md"
DOCX_PATH = DOCS_DIR / "reporte_arquitectura.docx"

BODY_FONT = "Times New Roman"
ACCENT = RGBColor(0x2C, 0x6E, 0x9C)

# ---------------------------------------------------------------------------
# Contenido (definido una vez; se renderiza a docx y a markdown)
# ---------------------------------------------------------------------------
TITLE = "Análisis de resultados — Arquitectura y producto final del sistema"

INTRO = (
    "El producto de este trabajo es un sistema de software funcional para la detección "
    "genotípica de resistencia a cefalosporinas de tercera generación en Escherichia coli. "
    "Tres piezas lo componen: un pipeline reproducible que encadena las herramientas "
    "bioinformáticas, una interfaz web para analizar una muestra suelta, y un contenedor que "
    "empaqueta todo para su distribución. Esta sección describe la arquitectura, las "
    "decisiones de diseño que sostienen su reproducibilidad, y los productos que el sistema "
    "entrega al usuario. Un caso real de ejecución —el aislado ERR17582235, descargado de "
    "ENA— sirve de hilo conductor."
)

SEC_ARCH_TITLE = "Arquitectura del sistema"
SEC_ARCH = (
    "El sistema sigue un diseño Python-first en el que cada lenguaje ocupa un rol acotado "
    "(Tabla 1). Python valida entradas, descarga y organiza archivos, invoca las herramientas "
    "externas, integra sus salidas y arma los reportes. Snakemake orquesta el flujo: cada "
    "regla declara sus entradas, salidas y comando, y el motor construye a partir de ellas un "
    "grafo dirigido acíclico que resuelve el orden de ejecución, paraleliza lo independiente y "
    "reanuda una corrida interrumpida sin repetir trabajo. R queda reservado para la "
    "estadística final. Ningún umbral vive dentro del código: la calidad mínima, los límites de "
    "ensamblaje, los cortes taxonómicos y de identidad se declaran en archivos de configuración "
    "YAML, de modo que ajustar el comportamiento no exige tocar los scripts."
)

SEC_FLOW_TITLE = "Flujo de procesamiento y herramientas"
SEC_FLOW = (
    "El pipeline encadena once etapas, cada una a cargo de una herramienta especializada y "
    "aislada en su propio ambiente (Tabla 2). El recorrido parte de las lecturas crudas y "
    "avanza por control de calidad, ensamblaje de novo, evaluación del ensamblaje, estimación "
    "de completitud, verificación de especie, detección de resistencia y tipificación, hasta "
    "converger en una tabla maestra y un reporte por muestra. La detección de resistencia se "
    "resuelve con dos motores independientes en lugar de uno, decisión de diseño que se retoma "
    "más abajo. La anotación estructural con Prokka existe como etapa disponible, pero no forma "
    "parte del recorrido por defecto."
)

SEC_QA_TITLE = "Decisiones de diseño para la reproducibilidad y la confianza"
SEC_QA_INTRO = (
    "Varias decisiones transversales distinguen al sistema de un conjunto de scripts encadenados "
    "y sostienen las propiedades que una tesis de este tipo debe demostrar."
)
SEC_QA_ITEMS = [
    ("Aislamiento de dependencias.", "Cada herramienta corre en su propio ambiente Conda con "
     "versiones fijadas. Se evita el conflicto real entre requisitos incompatibles (QUAST exige "
     "Python inferior a 3.12, por ejemplo) y se elimina la ambigüedad de \"qué versión produjo "
     "este resultado\"."),
    ("Orquestación determinista y reanudable.", "El grafo de dependencias de Snakemake fija el "
     "orden de ejecución a partir de los datos, no de un guion imperativo; una corrida "
     "interrumpida se retoma en el punto exacto en que quedó, sin rehacer lo ya calculado."),
    ("Trazabilidad.", "El sistema registra la versión de cada herramienta, además de hashes y "
     "fechas de los archivos descargados, de modo que cualquier resultado puede rastrearse hasta "
     "su origen y su entorno de cómputo."),
    ("Doble motor de detección.", "Los genes de resistencia se buscan con AMRFinderPlus y con "
     "ABricate de forma independiente; la concordancia entre ambos, resumida en el reporte, "
     "funciona como señal de alerta ante discrepancias en lugar de confiar en una sola fuente."),
    ("Controles de calidad explícitos.", "Cada etapa emite un estado PASS, WARNING o FAIL con "
     "umbrales configurables. Ninguna muestra se descarta en silencio: un fallo queda visible y "
     "acompañado de su motivo, y la verificación de especie con Kraken2 antecede a cualquier "
     "interpretación de resistencia."),
]

SEC_WEB_TITLE = "La interfaz de usuario"
SEC_WEB = (
    "La interfaz web es la vía de entrada para quien tiene una muestra suelta y quiere ver qué "
    "produce el sistema, sin editar tablas de configuración ni invocar Snakemake a mano "
    "(Figura 1). El formulario acepta dos tipos de entrada: un par de archivos FASTQ de lecturas "
    "crudas, que dispara el análisis completo, o un FASTA ya ensamblado, que corre un análisis "
    "parcial sin las etapas de calidad de lecturas, cobertura ni taxonomía dependiente de "
    "lecturas. Unos metadatos opcionales —plataforma de secuenciación, gen de resistencia "
    "esperado, número de hilos— afinan la ejecución y el reporte. Cada carga queda registrada "
    "como una corrida independiente y no interfiere con el lote curado principal."
)

SEC_REPORT_TITLE = "El reporte por muestra"
SEC_REPORT = (
    "El producto que el usuario recibe por cada muestra es un reporte HTML autocontenido "
    "(Figura 2). Sobre el aislado ERR17582235, el reporte encadena la identificación y "
    "procedencia; la calidad de las lecturas (2 086 206 lecturas iniciales, 50.97 % de GC); la "
    "cobertura estimada (61.35×, estado PASS); las métricas del ensamblaje (54 contigs, N50 de "
    "296 271 pb, PASS); la completitud y contaminación por CheckM (99.93 % y 0.26 %, PASS); la "
    "verificación taxonómica por Kraken2; la secuencia tipo (MLST); y los genes de resistencia "
    "detectados, presentados como tabla —con identidad, cobertura y clase de cada gen— y como "
    "gráfica. Sobre esta muestra, el pipeline identificó blaCMY-2 (β-lactamasa tipo AmpC) junto "
    "a determinantes de resistencia a quinolonas y fosfomicina. El reporte cierra con la "
    "interpretación del mecanismo, la comparación contra el estándar de referencia, la "
    "concordancia entre los dos motores de AMR y un bloque de advertencias en lenguaje llano. "
    "Cada valor viaja con su estado de control de calidad, y una nota destacada recuerda que el "
    "informe describe determinantes genotípicos, nunca una conclusión clínica."
)

SEC_DOCKER_TITLE = "Empaquetado y distribución"
SEC_DOCKER = (
    "Reproducir el sistema desde cero exige instalar doce herramientas con sus dependencias, un "
    "obstáculo real para quien solo quiere usarlo. El contenedor Docker resuelve ese costo: una "
    "sola imagen trae Snakemake, la interfaz web y los doce ambientes Conda ya construidos, de "
    "modo que un usuario con Docker levanta la interfaz o corre el pipeline con un puñado de "
    "comandos, sin instalar nada más. Las bases de datos de referencia grandes (Kraken2, CheckM) "
    "se montan como volumen para no inflar la imagen, mientras que la base compacta de "
    "AMRFinderPlus viaja dentro de ella. La imagen se construye para la arquitectura linux/amd64, "
    "lo que elimina la fricción de las herramientas de Bioconda en equipos con chip Apple."
)

SEC_PERF_TITLE = "Desempeño computacional"
SEC_PERF = (
    "Sobre las muestras de control, el procesamiento completo de un genoma tomó del orden de "
    "750 segundos, con un pico de memoria cercano a los 11 GB concentrado en el paso de "
    "colocación filogenética de CheckM. El tiempo de cómputo se mostró estable entre corridas "
    "repetidas (coeficiente de variación del 4 %), y el consumo de memoria pico algo más "
    "sensible a la carga del equipo (detalle en el reporte de validación estadística). La imagen "
    "del contenedor ocupa alrededor de 17.6 GB, dominada por los ambientes de las herramientas."
)

CLOSING = (
    "El resultado no es un análisis puntual sino un sistema reproducible y distribuible: los "
    "mismos datos y la misma configuración producen el mismo resultado, cada resultado es "
    "rastreable hasta su origen, y un tercero puede ejecutarlo sin reconstruir el entorno a mano."
)

# Tablas -------------------------------------------------------------------
TABLE1 = pd.DataFrame({
    "Componente": ["Python", "Snakemake", "R", "Conda"],
    "Rol en el sistema": [
        "Validación de entradas, descarga y organización, ejecución de herramientas, integración de resultados, generación de reportes y trazabilidad.",
        "Orquestación del flujo como grafo dirigido acíclico: dependencias, paralelización y reanudación.",
        "Análisis estadístico final (sensibilidad, especificidad, kappa, intervalos de confianza, coeficiente de variación).",
        "Aislamiento de dependencias: un ambiente por herramienta, con versiones fijadas.",
    ],
})

TABLE2 = pd.DataFrame({
    "Etapa": [
        "Descarga", "Control de calidad", "Ensamblaje", "Evaluación de ensamblaje",
        "Completitud / contaminación", "Verificación taxonómica", "Detección de resistencia",
        "Tipificación (MLST)", "Anotación (opcional)", "Estadística", "Reporte",
    ],
    "Herramienta": [
        "sra-tools", "fastp", "SPAdes", "QUAST", "CheckM", "Kraken2",
        "AMRFinderPlus + ABricate", "mlst", "Prokka", "R (caret, irr)", "Python (Jinja2)",
    ],
    "Función": [
        "Obtención de lecturas desde SRA/ENA.",
        "Recorte de adaptadores y filtrado por calidad y longitud.",
        "Ensamblaje de novo del genoma.",
        "Métricas del ensamblaje (contigs, N50, longitud).",
        "Integridad del genoma ensamblado.",
        "Confirmación de la especie antes de interpretar resistencia.",
        "Genes de resistencia, con dos motores independientes.",
        "Secuencia tipo, como contexto epidemiológico.",
        "Anotación estructural y funcional (fuera del flujo por defecto).",
        "Métricas de concordancia de la validación.",
        "Reporte HTML autocontenido por muestra.",
    ],
})


# ---------------------------------------------------------------------------
# Renderer Markdown
# ---------------------------------------------------------------------------
def df_to_md(df: pd.DataFrame) -> str:
    lines = ["| " + " | ".join(df.columns) + " |",
             "| " + " | ".join("---" for _ in df.columns) + " |"]
    for _, row in df.iterrows():
        lines.append("| " + " | ".join(str(v) for v in row) + " |")
    return "\n".join(lines)


def build_markdown() -> None:
    out = [f"# {TITLE}", "",
           "> Versión navegable del reporte en Word (`docs/reporte_arquitectura.docx`). "
           "Describe el sistema construido; complementa al reporte de "
           "[validación estadística](../validacion_estadistica/reporte/analisis_resultados.md), "
           "que analiza su concordancia analítica.", "",
           INTRO, "",
           f"## {SEC_ARCH_TITLE}", "", SEC_ARCH, "",
           "**Tabla 1.** Roles de los componentes del sistema.", "",
           df_to_md(TABLE1), "",
           f"## {SEC_FLOW_TITLE}", "", SEC_FLOW, "",
           "**Tabla 2.** Etapas del pipeline, herramienta responsable y función.", "",
           df_to_md(TABLE2), "",
           f"## {SEC_QA_TITLE}", "", SEC_QA_INTRO, ""]
    for head, body in SEC_QA_ITEMS:
        out.append(f"- **{head}** {body}")
    out += ["",
            f"## {SEC_WEB_TITLE}", "", SEC_WEB, "",
            "![Interfaz web del sistema](figuras/interfaz_web.png)", "",
            "*__Figura 1.__ Formulario de la interfaz web: identificador de la muestra, tipo de "
            "entrada (FASTQ pareado o FASTA ensamblado) y metadatos opcionales.*", "",
            f"## {SEC_REPORT_TITLE}", "", SEC_REPORT, "",
            "![Reporte por muestra](figuras/reporte_completo.png)", "",
            "*__Figura 2.__ Reporte HTML generado para el aislado ERR17582235, con sus secciones "
            "de calidad, ensamblaje, taxonomía, tipificación y detección de resistencia.*", "",
            f"## {SEC_DOCKER_TITLE}", "", SEC_DOCKER, "",
            f"## {SEC_PERF_TITLE}", "", SEC_PERF, "",
            "---", "", CLOSING, ""]
    MD_PATH.write_text("\n".join(out), encoding="utf-8")
    print(f"Markdown: {MD_PATH}")


# ---------------------------------------------------------------------------
# Renderer Word
# ---------------------------------------------------------------------------
def _set_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)


def _heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = BODY_FONT
        run.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x33, 0x33, 0x33)


def _body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _bullet(doc: Document, head: str, body: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(head + " ")
    r.bold = True
    p.add_run(body)


def _table_caption(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}. ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(10)


def _fig_caption(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}. ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(10)


def _table(doc: Document, df: pd.DataFrame, widths: list[float], font_size: int = 9) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, col in zip(table.rows[0].cells, df.columns):
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    for column, width in zip(table.columns, widths):
        for cell in column.cells:
            cell.width = Inches(width)


def _figure(doc: Document, image: Path, width: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Inches(width))


def build_docx() -> None:
    doc = Document()
    _set_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = ACCENT
    doc.add_paragraph()

    _heading(doc, "Análisis de resultados", 1)
    _body(doc, INTRO)

    _heading(doc, SEC_ARCH_TITLE, 2)
    _body(doc, SEC_ARCH)
    _table_caption(doc, "Tabla 1", "Roles de los componentes del sistema.")
    _table(doc, TABLE1, widths=[1.3, 5.0])

    _heading(doc, SEC_FLOW_TITLE, 2)
    _body(doc, SEC_FLOW)
    _table_caption(doc, "Tabla 2", "Etapas del pipeline, herramienta responsable y función.")
    _table(doc, TABLE2, widths=[1.7, 1.8, 3.0], font_size=8)

    _heading(doc, SEC_QA_TITLE, 2)
    _body(doc, SEC_QA_INTRO)
    for head, body in SEC_QA_ITEMS:
        _bullet(doc, head, body)

    _heading(doc, SEC_WEB_TITLE, 2)
    _body(doc, SEC_WEB)
    if (FIGURES_DIR / "interfaz_web.png").is_file():
        _figure(doc, FIGURES_DIR / "interfaz_web.png", width=5.6)
        _fig_caption(doc, "Figura 1", "Formulario de la interfaz web: identificador de la "
                     "muestra, tipo de entrada (FASTQ pareado o FASTA ensamblado) y metadatos "
                     "opcionales.")

    _heading(doc, SEC_REPORT_TITLE, 2)
    _body(doc, SEC_REPORT)
    if (FIGURES_DIR / "reporte_parte1.png").is_file():
        _figure(doc, FIGURES_DIR / "reporte_parte1.png", width=5.2)
        _figure(doc, FIGURES_DIR / "reporte_parte2.png", width=5.2)
        _fig_caption(doc, "Figura 2", "Reporte HTML generado para el aislado ERR17582235 "
                     "(dividido en dos partes por su extensión): identificación y controles de "
                     "calidad (arriba); detección de resistencia, interpretación y advertencias "
                     "(abajo).")

    _heading(doc, SEC_DOCKER_TITLE, 2)
    _body(doc, SEC_DOCKER)

    _heading(doc, SEC_PERF_TITLE, 2)
    _body(doc, SEC_PERF)
    _body(doc, CLOSING)

    doc.save(DOCX_PATH)
    print(f"Word: {DOCX_PATH}")


def main() -> None:
    build_markdown()
    build_docx()


if __name__ == "__main__":
    main()
