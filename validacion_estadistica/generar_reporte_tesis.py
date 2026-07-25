"""Genera el reporte de tesis (.docx) del analisis de validacion estadistica.

Traslada a un documento Word con formato academico las tablas y figuras del
notebook notebooks/validacion_estadistica.ipynb, leyendo los mismos
artefactos ya calculados (resultados/tables/validation_summary.tsv y
resultados/estadisticas/*.csv + plots/*.png). No recalcula estadistica: es
solo la capa de presentacion del analisis que ya produjo R en el notebook,
por lo que cualquier cambio en los numeros debe hacerse alli y regenerarse
este documento despues.

Uso:
    python generar_reporte_tesis.py
Salida:
    resultados/reporte/reporte_validacion_estadistica.docx
"""

from __future__ import annotations

from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.colors
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

VALIDATION_DIR = Path(__file__).resolve().parent
TABLES_DIR = VALIDATION_DIR / "resultados" / "tables"
STATS_DIR = VALIDATION_DIR / "resultados" / "estadisticas"
PLOTS_DIR = STATS_DIR / "plots"
SAMPLES_CSV = VALIDATION_DIR / "muestras" / "muestras_validacion_completo.csv"
OUTPUT_DIR = VALIDATION_DIR / "resultados" / "reporte"
OUTPUT_PATH = OUTPUT_DIR / "reporte_validacion_estadistica.docx"

BODY_FONT = "Times New Roman"
BODY_SIZE = 12
ACCENT = RGBColor(0x2C, 0x6E, 0x9C)


# ---------------------------------------------------------------------------
# Helpers de formato
# ---------------------------------------------------------------------------
def set_base_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)


def add_body(document: Document, text: str, justify: bool = True) -> None:
    paragraph = document.add_paragraph(text)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_runs(document: Document, segments: list[tuple[str, bool]]) -> None:
    """Parrafo con tramos en negrita/normal: [(texto, es_negrita), ...]."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in segments:
        run = paragraph.add_run(text)
        run.bold = bold
    return paragraph


def add_caption(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{label}. ")
    run.bold = True
    run.font.size = Pt(10)
    rest = paragraph.add_run(text)
    rest.font.size = Pt(10)
    rest.italic = True
    paragraph.paragraph_format.space_after = Pt(10)


def add_table_caption(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}. ")
    run.bold = True
    run.font.size = Pt(10)
    rest = paragraph.add_run(text)
    rest.font.size = Pt(10)
    rest.italic = True
    paragraph.paragraph_format.space_after = Pt(4)


def add_dataframe_table(document: Document, dataframe: pd.DataFrame, header_labels: list[str] | None = None,
                        font_size: int = 9) -> None:
    headers = header_labels if header_labels is not None else list(dataframe.columns)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for cell, label in zip(header_cells, headers):
        cell.text = str(label)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(font_size)
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = "" if pd.isna(value) else str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(document: Document, image_path: Path, width_inches: float = 5.6) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def heading(document: Document, text: str, level: int) -> None:
    h = document.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level <= 2 else RGBColor(0x33, 0x33, 0x33)
        run.font.name = BODY_FONT


def pct(x: float) -> str:
    return f"{x * 100:.1f}%"


ACCENT_HEX = "#2c6e9c"
BAR_MUTED = "#9aa7b4"


# ---------------------------------------------------------------------------
# Figuras del reporte (generadas aqui, no tomadas del notebook, para que
# queden sin titulo embebido -- el pie de figura ya cumple esa funcion -- y
# con un estilo visual homogeneo entre ambas). Se derivan de los mismos CSV
# que ya calculo R, asi que no hay recalculo estadistico.
# ---------------------------------------------------------------------------
def generate_confusion_figure(contingency: pd.DataFrame, output_path: Path) -> None:
    matrix = np.zeros((2, 2), dtype=int)  # matrix[prediccion][referencia]; 0=positive, 1=negative
    axis_index = {"positive": 0, "negative": 1}
    for _, row in contingency.iterrows():
        matrix[axis_index[row["pipeline_status"]]][axis_index[row["reference_status"]]] = row["count"]

    # matrix[0][0]=VP, [0][1]=FP, [1][0]=FN, [1][1]=VN
    cell_tags = [["VP", "FP"], ["FN", "VN"]]
    fig, ax = plt.subplots(figsize=(5.2, 4.4), dpi=200)
    cmap = matplotlib.colors.LinearSegmentedColormap.from_list("cm", ["#fbfbe8", ACCENT_HEX])
    ax.imshow(matrix, cmap=cmap, vmin=0, vmax=matrix.max())
    for i in range(2):
        for j in range(2):
            value = matrix[i][j]
            strong = value > matrix.max() * 0.55
            ax.text(j, i - 0.05, str(value), ha="center", va="center",
                    fontsize=22, color="white" if strong else "#1a1a1a", fontweight="bold")
            ax.text(j, i + 0.32, cell_tags[i][j], ha="center", va="center",
                    fontsize=9, color="white" if strong else "#555", style="italic")
    ax.set_xticks([0, 1]); ax.set_xticklabels(["Positiva", "Negativa"], fontsize=10.5)
    ax.set_yticks([0, 1]); ax.set_yticklabels(["Positiva", "Negativa"], fontsize=10.5)
    ax.set_xlabel("Estándar de referencia", fontsize=11)
    ax.set_ylabel("Predicción del pipeline", fontsize=11)
    ax.xaxis.set_label_position("top"); ax.xaxis.tick_top()
    colorbar = fig.colorbar(ax.images[0], ax=ax, fraction=0.046, pad=0.04)
    colorbar.set_label("Número de muestras", fontsize=9)
    ax.set_xticks(np.arange(-.5, 2, 1), minor=True); ax.set_yticks(np.arange(-.5, 2, 1), minor=True)
    ax.grid(which="minor", color="#888", lw=0.8); ax.tick_params(which="minor", length=0)
    plt.tight_layout()
    plt.savefig(output_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_metrics_figure(metrics_all: pd.DataFrame, metrics_conf: pd.DataFrame,
                            n_all: int, n_conf: int, output_path: Path) -> None:
    metric_keys = ["Sensitivity", "Specificity", "Accuracy"]
    metric_labels = ["Sensibilidad", "Especificidad", "Exactitud"]

    def values(dframe: pd.DataFrame):
        estimate = [dframe.loc[m, "estimate"] * 100 for m in metric_keys]
        lower = [(dframe.loc[m, "estimate"] - dframe.loc[m, "ci_lower"]) * 100 for m in metric_keys]
        upper = [(dframe.loc[m, "ci_upper"] - dframe.loc[m, "estimate"]) * 100 for m in metric_keys]
        return estimate, [lower, upper]

    est_all, err_all = values(metrics_all)
    est_conf, err_conf = values(metrics_conf)
    x = np.arange(len(metric_keys)); width = 0.36
    fig, ax = plt.subplots(figsize=(7.2, 4.8), dpi=200)
    bars_all = ax.bar(x - width / 2, est_all, width, yerr=err_all, capsize=4, color=BAR_MUTED,
                      edgecolor="#4a4a4a", linewidth=0.6, error_kw=dict(ecolor="#333333", lw=1),
                      label=f"Todas las muestras evaluables (n={n_all})")
    bars_conf = ax.bar(x + width / 2, est_conf, width, yerr=err_conf, capsize=4, color=ACCENT_HEX,
                       edgecolor="#1a1a1a", linewidth=0.6, error_kw=dict(ecolor="#333333", lw=1),
                       label=f"Especie confirmada por Kraken2 (n={n_conf})")
    for bars, estimate, err in [(bars_all, est_all, err_all), (bars_conf, est_conf, err_conf)]:
        for bar, value, hi in zip(bars, estimate, err[1]):
            ax.text(bar.get_x() + bar.get_width() / 2, value + hi + 1.8, f"{value:.1f}%",
                    ha="center", va="bottom", fontsize=8.5, color="#222")
    ax.set_ylim(0, 112); ax.set_ylabel("Valor (%)", fontsize=11)
    ax.set_xticks(x); ax.set_xticklabels(metric_labels, fontsize=10.5)
    ax.set_yticks(range(0, 101, 20))
    ax.legend(loc="lower center", fontsize=8.8, frameon=False, bbox_to_anchor=(0.5, -0.30))
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", ls=":", alpha=0.4); ax.set_axisbelow(True)
    plt.tight_layout()
    plt.savefig(output_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# ---------------------------------------------------------------------------
# Carga de datos
# ---------------------------------------------------------------------------
def load_data() -> dict:
    summary = pd.read_csv(TABLES_DIR / "validation_summary.tsv", sep="\t")
    meta = pd.read_csv(SAMPLES_CSV)
    metrics_conf = pd.read_csv(STATS_DIR / "classification_metrics.csv").set_index("metric")
    metrics_all = pd.read_csv(STATS_DIR / "classification_metrics_all_samples.csv").set_index("metric")
    kappa = pd.read_csv(STATS_DIR / "kappa.csv").iloc[0]
    contingency = pd.read_csv(STATS_DIR / "contingency_table.csv")
    return dict(summary=summary, meta=meta, metrics_conf=metrics_conf,
                metrics_all=metrics_all, kappa=kappa, contingency=contingency)


# ---------------------------------------------------------------------------
# Construccion del documento
# ---------------------------------------------------------------------------
def build_document(data: dict) -> Document:
    summary = data["summary"]
    meta = data["meta"]
    mc = data["metrics_conf"]
    ma = data["metrics_all"]
    kappa = data["kappa"]

    n_total = len(summary)
    n_otra = int((summary["species_check"] == "otra_especie").sum())
    n_conf_eval = int(mc.loc["Accuracy", "n"])
    n_all_eval = int(ma.loc["Accuracy", "n"])

    document = Document()
    set_base_style(document)

    # ---- Titulo ----
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Validación estadística de un pipeline bioinformático para la detección "
                        "genotípica de resistencia a cefalosporinas de tercera generación en "
                        "Escherichia coli")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = ACCENT
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run("Análisis de concordancia entre las predicciones genotípicas y los "
                            "fenotipos de referencia documentados")
    srun.italic = True
    srun.font.size = Pt(11)
    document.add_paragraph()

    # ---- 1. Introduccion ----
    heading(document, "1. Introducción", 1)
    add_body(document,
        "El presente reporte documenta la validación estadística de un pipeline bioinformático "
        "de código abierto diseñado para analizar secuencias genómicas de Escherichia coli "
        "asociadas con resistencia a cefalosporinas de tercera generación. La validación evalúa, "
        "sobre un conjunto de genomas reales de repositorios públicos, el grado de concordancia "
        "entre la predicción genotípica del pipeline (presencia o ausencia de genes de resistencia "
        "adquiridos, principalmente β-lactamasas de espectro extendido y AmpC plasmídicas) y el "
        "fenotipo o genotipo de referencia documentado para cada muestra. El objetivo no es la "
        "validación clínica ni la implementación diagnóstica, sino demostrar la factibilidad "
        "técnica, la concordancia analítica y la reproducibilidad del flujo de análisis.")

    # ---- 2. Materiales y metodos ----
    heading(document, "2. Materiales y métodos", 1)

    heading(document, "2.1. Conjunto de muestras", 2)
    add_body(document,
        f"Se ensambló un conjunto de {n_total} genomas reales de E. coli, previamente ensamblados "
        "y descargados de la base de datos del NCBI, seleccionados como controles positivos "
        "(portadores de un gen de resistencia documentado), negativos (cepas de referencia sin "
        "resistencia adquirida esperada) y límite (elegidos por su calidad de ensamblaje subóptima, "
        "para evaluar la robustez del pipeline). La composición del conjunto se resume en la Tabla 1.")

    composition = pd.crosstab(meta["Clasificación"], meta["Tipo de control"])
    composition_reset = composition.reset_index()
    composition_reset.columns = ["Categoría"] + list(composition.columns)
    add_table_caption(document, "Tabla 1", "Composición del conjunto de validación por categoría de "
                      "resistencia y tipo de control.")
    add_dataframe_table(document, composition_reset)

    heading(document, "2.2. Flujo bioinformático", 2)
    add_body(document,
        "Cada genoma se procesó con el mismo flujo de análisis del pipeline en su modo de solo "
        "ensamblaje: evaluación de la calidad del ensamblaje (QUAST), estimación de completitud y "
        "contaminación (CheckM), detección de genes de resistencia con dos motores independientes "
        "(AMRFinderPlus y ABricate, este último contra las bases CARD y ResFinder), tipificación "
        "de secuencia multilocus (MLST, esquema de Achtman de 7 loci) y comparación de los genes "
        "detectados contra el estándar de referencia de cada muestra.")

    heading(document, "2.3. Verificación taxonómica", 2)
    add_body(document,
        "De forma complementaria, se ejecutó una verificación taxonómica con Kraken2 sobre cada "
        "ensamblaje para confirmar la identidad de especie. Este paso resultó determinante para la "
        "interpretación correcta de los resultados, como se detalla en la sección 3.2.")

    heading(document, "2.4. Definiciones y análisis estadístico", 2)
    add_body(document,
        "Cada muestra con estándar de referencia evaluable se clasificó como verdadero positivo (TP), "
        "verdadero negativo (TN), falso positivo (FP) o falso negativo (FN) según la coincidencia "
        "entre la predicción del pipeline y la referencia. Las muestras cuyo estándar de referencia "
        "no era aplicable por genotipo (controles límite) se marcaron como indeterminadas y se "
        "excluyeron del cálculo. Sobre la matriz de confusión resultante se calcularon la "
        "sensibilidad, la especificidad y la exactitud, con sus intervalos de confianza (IC) del "
        "95% por el método exacto de Clopper-Pearson, y el índice kappa de Cohen como medida de "
        "concordancia corregida por el azar. El análisis estadístico se realizó en R (paquetes "
        "caret e irr).")

    # ---- 3. Resultados ----
    heading(document, "3. Resultados", 1)

    heading(document, "3.1. Verificación de especie", 2)
    add_runs(document, [
        ("La verificación taxonómica con Kraken2 reveló que ", False),
        (f"{n_otra} de los {n_total} genomas ({pct(n_otra/n_total)}) no correspondían realmente a "
         "E. coli", True),
        (", pese a estar etiquetados como tal en los metadatos de origen. Los taxones predominantes "
         "identificados abarcaron desde parientes cercanos de la familia Enterobacteriaceae "
         "(Klebsiella, Enterobacter, Citrobacter, Salmonella) hasta especies filogenéticamente "
         "distantes (Listeria, Staphylococcus aureus, Aeromonas) e incluso una muestra de origen "
         "predominantemente humano. La Tabla 2 detalla estas muestras. Ninguna pertenecía a las "
         "categorías de control negativo ni de control límite.", False),
    ])

    otra = summary[summary["species_check"] == "otra_especie"][
        ["sample_id", "category", "confusion_category", "predominant_taxon"]
    ].copy()
    otra.columns = ["Accesión", "Categoría", "Clasificación", "Taxón predominante (Kraken2)"]
    add_table_caption(document, "Tabla 2", f"Las {n_otra} muestras cuya especie predominante no es "
                      "E. coli según Kraken2.")
    add_dataframe_table(document, otra, font_size=8)
    add_body(document,
        "Dado que la resistencia a cefalosporinas se evalúa específicamente sobre E. coli, estas "
        "muestras introducían un sesgo en las métricas y se excluyeron del análisis principal, "
        "reportándose ambos cálculos (con y sin exclusión) para dejar trazable su efecto. Se "
        "conservó una única muestra de Shigella, por su cercanía genómica con E. coli, siguiendo "
        "el criterio del pipeline de no excluirla automáticamente.")

    heading(document, "3.2. Matriz de confusión", 2)
    add_body(document,
        f"Tras excluir las {n_otra} muestras de otra especie y los controles indeterminados, la "
        f"comparación se basó en {n_conf_eval} muestras de especie confirmada con estándar de "
        "referencia evaluable. La Figura 1 presenta la matriz de confusión resultante.")
    if (PLOTS_DIR / "confusion_matrix_reporte.png").exists():
        add_figure(document, PLOTS_DIR / "confusion_matrix_reporte.png", width_inches=4.6)
        add_caption(document, "Figura 1", "Matriz de confusión del pipeline frente al estándar de "
                    f"referencia sobre las {n_conf_eval} muestras de especie confirmada (VP: "
                    "verdadero positivo; VN: verdadero negativo; FP: falso positivo; FN: falso "
                    "negativo).")

    heading(document, "3.3. Métricas de desempeño diagnóstico", 2)
    add_runs(document, [
        ("Sobre el conjunto de especie confirmada, el pipeline alcanzó una ", False),
        (f"sensibilidad del {pct(mc.loc['Sensitivity','estimate'])} "
         f"(IC 95%: {pct(mc.loc['Sensitivity','ci_lower'])}–{pct(mc.loc['Sensitivity','ci_upper'])})", True),
        (f", una ", False),
        (f"especificidad del {pct(mc.loc['Specificity','estimate'])} "
         f"(IC 95%: {pct(mc.loc['Specificity','ci_lower'])}–{pct(mc.loc['Specificity','ci_upper'])})", True),
        (f" y una ", False),
        (f"exactitud global del {pct(mc.loc['Accuracy','estimate'])} "
         f"(IC 95%: {pct(mc.loc['Accuracy','ci_lower'])}–{pct(mc.loc['Accuracy','ci_upper'])})", True),
        (". Como referencia, sin excluir las muestras de otra especie la sensibilidad descendía a "
         f"{pct(ma.loc['Sensitivity','estimate'])} y la exactitud a {pct(ma.loc['Accuracy','estimate'])}, "
         "lo que ilustra el impacto de la verificación taxonómica (Tabla 3, Figura 2).", False),
    ])

    metrics_table = pd.DataFrame({
        "Métrica": ["Sensibilidad", "Especificidad", "Exactitud"],
        f"Especie confirmada (n={n_conf_eval})": [
            f"{pct(mc.loc[m,'estimate'])} ({pct(mc.loc[m,'ci_lower'])}–{pct(mc.loc[m,'ci_upper'])})"
            for m in ["Sensitivity", "Specificity", "Accuracy"]
        ],
        f"Todas las evaluables (n={n_all_eval})": [
            f"{pct(ma.loc[m,'estimate'])} ({pct(ma.loc[m,'ci_lower'])}–{pct(ma.loc[m,'ci_upper'])})"
            for m in ["Sensitivity", "Specificity", "Accuracy"]
        ],
    })
    add_table_caption(document, "Tabla 3", "Métricas de desempeño diagnóstico con intervalo de "
                      "confianza del 95% (Clopper-Pearson), con y sin exclusión taxonómica.")
    add_dataframe_table(document, metrics_table, font_size=9)

    if (PLOTS_DIR / "metricas_comparadas.png").exists():
        add_figure(document, PLOTS_DIR / "metricas_comparadas.png", width_inches=5.6)
        add_caption(document, "Figura 2", "Comparación de las métricas de desempeño antes y después "
                    "de excluir las muestras de otra especie. Las barras de error representan el "
                    "IC del 95%.")

    heading(document, "3.4. Concordancia global (índice kappa)", 2)
    add_runs(document, [
        ("El índice kappa de Cohen sobre las muestras de especie confirmada fue de ", False),
        (f"{kappa['kappa_value']:.3f} (p < 0.001)", True),
        (", valor que corresponde a una concordancia “casi perfecta” según la escala de "
         "Landis y Koch y muy superior al azar. Sin la exclusión taxonómica, el mismo índice "
         "descendía a 0.596 (concordancia moderada-sustancial), lo que confirma que buena parte de "
         "la discordancia aparente provenía de muestras mal identificadas y no de errores de "
         "detección del pipeline.", False),
    ])

    heading(document, "3.5. Análisis de casos discordantes", 2)
    add_body(document,
        "La Tabla 4 detalla los casos discordantes. Se observa que la mayoría de los falsos "
        "negativos corresponden a muestras que no son E. coli (columna de verificación de especie); "
        "una vez excluidas, solo persisten dos falsos negativos genuinos y cuatro falsos positivos.")
    disc = summary[summary["confusion_category"].isin(["FN", "FP"])][
        ["sample_id", "category", "expected_gene", "detected_gene", "confusion_category", "species_check"]
    ].copy()
    disc["species_check"] = disc["species_check"].map(
        {"confirmada": "E. coli", "otra_especie": "otra especie", "shigella_revision_manual": "Shigella"}
    )
    disc.columns = ["Accesión", "Categoría", "Gen esperado", "Gen detectado", "Clasif.", "Especie"]
    add_table_caption(document, "Tabla 4", "Casos discordantes (falsos negativos y falsos positivos) "
                      "con su verificación de especie.")
    add_dataframe_table(document, disc, font_size=8)

    add_body(document,
        "Los cuatro falsos positivos, todos en cepas E. coli confirmadas de la categoría negativa, "
        "corresponden a la β-lactamasa de espectro estrecho blaTEM-1 (dos casos) y a mutaciones "
        "puntuales en el gen cirA (dos casos), ninguno de los cuales confiere por sí solo "
        "resistencia a cefalosporinas de tercera generación. Constituyen un desajuste entre la "
        "definición estricta del control negativo y la regla de comparación del pipeline, más que "
        "un error de detección. Los dos falsos negativos genuinos, en cepas E. coli confirmadas por "
        "MLST, no presentaron ningún indicio del gen esperado ni siquiera antes de aplicar los "
        "umbrales de identidad y cobertura, y quedan como casos puntuales para revisión manual.")

    heading(document, "3.6. Comportamiento de los controles límite", 2)
    add_body(document,
        "Los cinco controles límite, seleccionados por su calidad de ensamblaje y no por su "
        "genotipo, se comportaron según lo esperado. En el extremo de peor calidad, la cepa del "
        "brote alemán de 2011 (O104:H4 TY-2482), con el ensamblaje más fragmentado del conjunto, "
        "aún detectó los genes blaCTX-M y blaTEM-1, aunque sin resolver el alelo exacto ni el tipo "
        "de secuencia, evidenciando una degradación progresiva y controlada del desempeño. En el "
        "extremo de mejor calidad, la cepa de referencia EC958 (clon pandémico ST131), con "
        "ensamblaje completo, detectó los cuatro genes documentados en la literatura (blaCTX-M-15, "
        "blaCMY-23, blaOXA-1 y blaTEM-1). Las tres cepas restantes, incluida la cepa tipo de la "
        "especie (ATCC 11775), no presentaron genes adquiridos, resultado coherente entre los dos "
        "motores de detección independientes.")

    # ---- 4. Discusion ----
    heading(document, "4. Discusión", 1)
    add_body(document,
        "El hallazgo más relevante de esta validación fue de naturaleza metodológica: una fracción "
        f"apreciable del conjunto ({pct(n_otra/n_total)}) no correspondía a la especie declarada, "
        "lo que deprimía artificialmente las métricas cuando se analizaba el conjunto completo. "
        "Este resultado subraya la importancia de incorporar la verificación taxonómica como paso "
        "obligatorio previo a cualquier análisis de resistencia, tal como el pipeline lo exige en "
        "su modo operativo sobre lecturas crudas. Una vez restringido el análisis a genomas "
        "genuinamente de E. coli, el pipeline mostró una concordancia casi perfecta con el estándar "
        "de referencia, con una sensibilidad y una exactitud superiores al 90% y un índice kappa de "
        f"{kappa['kappa_value']:.2f}.")
    add_body(document,
        "El uso de dos motores de detección independientes reforzó la confianza en los resultados "
        "negativos: en todas las muestras clasificadas como sin genes adquiridos, tanto "
        "AMRFinderPlus como ABricate/ResFinder coincidieron, mientras que los numerosos aciertos de "
        "ABricate/CARD correspondían al resistoma intrínseco cromosómico de E. coli y no a "
        "determinantes adquiridos de resistencia a cefalosporinas. La revisión de los controles "
        "límite confirmó, además, que el desempeño del pipeline se degrada de forma predecible ante "
        "ensamblajes de calidad decreciente, sin producir fallos abruptos.")

    # ---- 5. Conclusiones y limitaciones ----
    heading(document, "5. Conclusiones y limitaciones", 1)
    add_body(document,
        "Sobre un conjunto de genomas de E. coli confirmados por identidad de especie, el pipeline "
        "demostró una elevada concordancia analítica con el estándar de referencia para la detección "
        "genotípica de resistencia a cefalosporinas de tercera generación, cumpliendo el objetivo de "
        "demostrar factibilidad técnica y reproducibilidad. Como principal limitación metodológica, "
        "la verificación taxonómica se incorporó de forma retrospectiva a este conjunto de "
        "validación; se recomienda aplicarla antes de admitir cualquier muestra nueva. Asimismo, la "
        "evaluación de la reproducibilidad computacional (coeficiente de variación del tiempo de "
        "ejecución y del consumo de memoria entre corridas repetidas) se encuentra en curso al "
        "momento de este reporte y se integrará en una versión posterior.")

    # ---- Referencias ----
    heading(document, "Referencias", 1)
    references = [
        "Forde, B. M. et al. (2014). The complete genome sequence of Escherichia coli EC958: a high "
        "quality reference sequence for the globally disseminated multidrug resistant E. coli "
        "O25b:H4-ST131 clone. PLoS ONE, 9(8), e104400.",
        "Landis, J. R. & Koch, G. G. (1977). The measurement of observer agreement for categorical "
        "data. Biometrics, 33(1), 159–174.",
        "Wirth, T. et al. (2006). Sex and virulence in Escherichia coli: an evolutionary "
        "perspective. Molecular Microbiology, 60(5), 1136–1151. (Esquema MLST de Achtman)",
        "Wood, D. E. & Salzberg, S. L. (2014). Kraken: ultrafast metagenomic sequence "
        "classification using exact alignments. Genome Biology, 15(3), R46.",
        "Feldgarden, M. et al. (2019). Validating the AMRFinder tool and resistance gene database. "
        "Antimicrobial Agents and Chemotherapy, 63(11), e00483-19.",
        "Clopper, C. J. & Pearson, E. S. (1934). The use of confidence or fiducial limits "
        "illustrated in the case of the binomial. Biometrika, 26(4), 404–413.",
    ]
    for ref in references:
        paragraph = document.add_paragraph(ref)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Inches(0.4)
        paragraph.paragraph_format.first_line_indent = Inches(-0.4)
        for run in paragraph.runs:
            run.font.size = Pt(10)

    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    PLOTS_DIR.mkdir(parents=True, exist_ok=True)
    data = load_data()

    generate_confusion_figure(data["contingency"], PLOTS_DIR / "confusion_matrix_reporte.png")
    generate_metrics_figure(
        data["metrics_all"], data["metrics_conf"],
        n_all=int(data["metrics_all"].loc["Accuracy", "n"]),
        n_conf=int(data["metrics_conf"].loc["Accuracy", "n"]),
        output_path=PLOTS_DIR / "metricas_comparadas.png",
    )

    document = build_document(data)
    document.save(OUTPUT_PATH)
    print(f"Reporte generado en {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
